# In core/views.py

import os
import re
import tempfile
import traceback
import spacy

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.decorators import login_required
from django.views.generic import FormView
from django.http import JsonResponse, HttpResponseBadRequest
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from django.contrib import messages
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Q, Prefetch
from django.views.decorators.http import require_POST

from .forms import SignUpForm, LoginForm, FreelancerDataForm, RecruiterDataForm, JobPostForm
from .models import FreelancerData, Application, RecruiterData, Job, Skill

# --- Recommendation Engine Imports ---
try:
    from recommendations.recommender import get_job_recommendations, rank_applications, get_resume_ats_score
except ImportError:
    def get_job_recommendations(freelancer): return []
    def rank_applications(job): return Application.objects.filter(job=job)
    def get_resume_ats_score(job_text, resume_text): return 0

# --- Resume Parsing Imports ---
try:
    from pdfminer.high_level import extract_text as extract_pdf_text
except ImportError:
    extract_pdf_text = None

try:
    import docx
except ImportError:
    print("Warning: python-docx not installed. DOCX parsing disabled.")
    docx = None

try:
    from PIL import Image
    import pytesseract
    # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
except ImportError:
    print("Warning: Pillow or pytesseract not installed. OCR disabled.")
    Image = None
    pytesseract = None
except FileNotFoundError:
    print("ERROR: Tesseract executable not found. OCR disabled.")
    Image = None
    pytesseract = None

# --- spaCy Model Loading ---
try:
    nlp = spacy.load('en_core_web_sm')
    print("Default spaCy model 'en_core_web_sm' loaded.")
except (ImportError, OSError):
    nlp = None
    print("Warning: Default spaCy model not found.")

try:
    custom_nlp = spacy.load("./custom_ner_model")
    print("Custom NER model for skills loaded successfully.")
except (IOError):
    custom_nlp = None
    print("Warning: Custom NER model not found.")


# ==============================================================================
# CORE & AUTHENTICATION VIEWS
# ==============================================================================

def home(request):
    return render(request, "core/home.html")

class SignUpView(FormView):
    template_name = 'core/signup.html'
    form_class = SignUpForm

    def form_valid(self, form):
        user = form.save()
        login(self.request, user)
        messages.success(self.request, "Registration successful! Please complete your profile.")
        user_type = form.cleaned_data['type']
        if user_type == 'freelancer':
            return redirect('edit_profile')
        else:
            return redirect('recruiter_profile_edit')

class LoginView(FormView):
    template_name = 'core/login.html'
    form_class = LoginForm

    def form_valid(self, form):
        user = form.get_user()
        login(self.request, user)
        messages.success(self.request, f"Welcome back, {user.username}!")
        if user.is_freelancer:
            return redirect('freelancer_dashboard')
        elif user.is_recruiter:
            return redirect('recruiter_dashboard')
        return redirect('home')

def logout_view(request):
    logout(request)
    messages.success(request, "You have been successfully logged out.")
    return redirect('home')

@login_required
@require_POST
def delete_account_view(request):
    user = request.user
    username_deleted = user.username
    logout(request)
    user.delete()
    messages.success(request, f"Account '{username_deleted}' has been permanently deleted.")
    return redirect('home')

# ==============================================================================
# FREELANCER VIEWS
# ==============================================================================

@login_required
def freelancer_dashboard(request):
    try:
        freelancer_instance = FreelancerData.objects.get(user=request.user)
    except FreelancerData.DoesNotExist:
        messages.warning(request, "Please complete your profile to access the dashboard.")
        return redirect('edit_profile')

    accepted_applications = Application.objects.filter(
        freelancer=freelancer_instance,
        status='ACCEPTED'
    ).select_related('job', 'job__recruiter').order_by('-job__posted_at')

    recommended_jobs = get_job_recommendations(freelancer_instance)

    context = {
        'freelancer': freelancer_instance,
        'accepted_gigs': accepted_applications,
        'recommended_jobs': recommended_jobs,
    }
    return render(request, 'core/f_dashboard.html', context)

@login_required
def create_freelancer_profile(request):
    try:
        freelancer_instance = FreelancerData.objects.get(user=request.user)
        is_editing = True
    except FreelancerData.DoesNotExist:
        freelancer_instance = None
        is_editing = False

    if request.method == 'POST':
        form = FreelancerDataForm(request.POST, request.FILES, instance=freelancer_instance)
        if form.is_valid():
            profile = form.save(commit=False)
            profile.user = request.user
            profile.save()
            form.save_m2m()
            messages.success(request, f"Profile {'updated' if is_editing else 'created'} successfully!")
            return redirect('freelancer_dashboard')
    else:
        form = FreelancerDataForm(instance=freelancer_instance)

    context = {
        'form': form,
        'is_editing': is_editing,
        'current_skills': freelancer_instance.skills.all() if freelancer_instance else Skill.objects.none()
    }
    return render(request, 'core/freelanceredit_profile.html', context)

@login_required
@require_POST
def remove_freelancer_skill(request):
    try:
        freelancer_profile = FreelancerData.objects.get(user=request.user)
        skill_id = int(request.POST.get('skill_id'))
        skill_to_remove = get_object_or_404(Skill, pk=skill_id)
        freelancer_profile.skills.remove(skill_to_remove)
        return JsonResponse({'success': True, 'message': 'Skill removed.'})
    except FreelancerData.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Freelancer profile not found.'}, status=404)
    except (ValueError, TypeError):
        return JsonResponse({'success': False, 'error': 'Invalid Skill ID.'}, status=400)
    except Exception:
        return JsonResponse({'success': False, 'error': 'An unexpected error occurred.'}, status=500)


# ==============================================================================
# RECRUITER & JOB VIEWS
# ==============================================================================

@login_required
def recruiter_dashboard(request):
    try:
        recruiter = RecruiterData.objects.get(user=request.user)
    except RecruiterData.DoesNotExist:
        messages.warning(request, "Please complete your recruiter profile.")
        return redirect('recruiter_profile_edit')
    context = {'recruiter': recruiter}
    return render(request, 'core/r_dashboard.html', context)

@login_required
def recruiter_profile_edit(request):
    try:
        recruiter_instance = RecruiterData.objects.get(user=request.user)
        is_editing = True
    except RecruiterData.DoesNotExist:
        recruiter_instance = None
        is_editing = False

    if request.method == 'POST':
        form = RecruiterDataForm(request.POST, request.FILES, instance=recruiter_instance)
        if form.is_valid():
            profile = form.save(commit=False)
            profile.user = request.user
            profile.save()
            messages.success(request, f"Profile {'updated' if is_editing else 'created'} successfully!")
            return redirect('recruiter_dashboard')
    else:
        form = RecruiterDataForm(instance=recruiter_instance)

    context = {'form': form, 'is_editing': is_editing}
    return render(request, 'core/recruiteredit_profile.html', context)

@login_required
def manage_job(request, job_id=None):
    recruiter_instance = get_object_or_404(RecruiterData, user=request.user)
    job_instance = None
    if job_id:
        job_instance = get_object_or_404(Job, pk=job_id, recruiter=recruiter_instance)

    if request.method == 'POST':
        form = JobPostForm(request.POST, instance=job_instance)
        if form.is_valid():
            job = form.save(commit=False)
            job.recruiter = recruiter_instance
            job.save()
            form.save_m2m()
            messages.success(request, f"Job {'updated' if job_instance else 'posted'} successfully!")
            return redirect('recruiter_job_list')
    else:
        form = JobPostForm(instance=job_instance)

    context = {'form': form, 'is_editing': job_instance is not None}
    return render(request, 'core/jobs_form.html', context)

@login_required
def recruited_job_list(request):
    recruiter_instance = get_object_or_404(RecruiterData, user=request.user)
    jobs = Job.objects.filter(recruiter=recruiter_instance).prefetch_related(
        Prefetch('applications', queryset=Application.objects.filter(status='ACCEPTED'), to_attr='accepted_application_list')
    )
    context = {'jobs': jobs}
    return render(request, 'core/recruiter_job_list.html', context)


# ==============================================================================
# PUBLIC JOB & APPLICATION VIEWS
# ==============================================================================

def job_list(request):
    queryset = Job.objects.filter(is_active=True).select_related('recruiter').prefetch_related('required_skills')
    query = request.GET.get('q', '').strip()
    location_query = request.GET.get('location', '').strip()
    skill_ids = request.GET.getlist('skills')

    if query:
        queryset = queryset.filter(Q(title__icontains=query) | Q(description__icontains=query))
    if location_query:
        queryset = queryset.filter(location__icontains=location_query)
    if skill_ids:
        queryset = queryset.filter(required_skills__id__in=skill_ids).distinct()

    paginator = Paginator(queryset, 10)
    page_number = request.GET.get('page')
    jobs_page = paginator.get_page(page_number)

    context = {
        'jobs': jobs_page,
        'all_skills': Skill.objects.all().order_by('name'),
        'current_query': query,
        'current_location': location_query,
        'current_skills': skill_ids,
    }
    return render(request, 'core/job_list.html', context)

def job_detail(request, job_id):
    job = get_object_or_404(Job.objects.select_related('recruiter'), pk=job_id, is_active=True)
    has_applied = False
    is_freelancer = False
    if request.user.is_authenticated and hasattr(request.user, 'freelancer_profile'):
        is_freelancer = True
        has_applied = Application.objects.filter(job=job, freelancer=request.user.freelancer_profile).exists()

    context = {'job': job, 'has_applied': has_applied, 'is_freelancer': is_freelancer}
    return render(request, 'core/job_detail.html', context)

@login_required
@require_POST
def apply_to_job(request, job_id):
    job = get_object_or_404(Job, pk=job_id, is_active=True)
    freelancer = get_object_or_404(FreelancerData, user=request.user)

    if Application.objects.filter(job=job, freelancer=freelancer).exists():
        messages.warning(request, "You have already applied for this job.")
        return redirect('job_detail', job_id=job.id)

    Application.objects.create(
        job=job, freelancer=freelancer, cover_letter=request.POST.get('cover_letter', '').strip()
    )
    messages.success(request, f"Successfully applied for the job: {job.title}")
    return redirect('job_detail', job_id=job.id)

@login_required
def view_job_applications(request, job_id):
    recruiter = get_object_or_404(RecruiterData, user=request.user)
    job = get_object_or_404(Job, pk=job_id, recruiter=recruiter)
    ranked_applications = rank_applications(job)
    context = {'job': job, 'applications': ranked_applications}
    return render(request, 'core/job_applications.html', context)

@login_required
@require_POST
def update_application_status(request, application_id):
    recruiter = get_object_or_404(RecruiterData, user=request.user)
    application = get_object_or_404(Application, pk=application_id, job__recruiter=recruiter)
    new_status = request.POST.get('status')
    if new_status in ['ACCEPTED', 'DECLINED']:
        application.status = new_status
        application.save(update_fields=['status'])
        messages.success(request, f"Application status updated to {application.get_status_display()}.")
    else:
        messages.error(request, "Invalid status update requested.")
    return redirect('view_job_applications', job_id=application.job.id)

# ==============================================================================
# AJAX & UTILITY VIEWS
# ==============================================================================

@login_required
@require_POST
def parse_resume_view(request):
    uploaded_file = request.FILES.get('resume_file')
    if not uploaded_file:
        return JsonResponse({'error': 'No resume file provided.'}, status=400)

    # File Validation
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    if file_extension not in ['.pdf', '.docx', '.jpg', '.jpeg', '.png']:
        return JsonResponse({'error': 'Invalid file type.'}, status=400)
    if uploaded_file.size > 5 * 1024 * 1024:
        return JsonResponse({'error': 'File size exceeds 5MB.'}, status=400)

    # Text Extraction
    extracted_text = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            for chunk in uploaded_file.chunks(): temp_file.write(chunk)
            temp_file_path = temp_file.name
        
        if file_extension == '.pdf' and extract_pdf_text:
            extracted_text = extract_pdf_text(temp_file_path)
        elif file_extension == '.docx' and docx:
            extracted_text = '\n'.join([p.text for p in docx.Document(temp_file_path).paragraphs])
        elif file_extension in ['.jpg', '.jpeg', '.png'] and Image and pytesseract:
            extracted_text = pytesseract.image_to_string(Image.open(temp_file_path))
        
        os.unlink(temp_file_path)
    except Exception:
        traceback.print_exc()
        return JsonResponse({'error': 'Failed to read or process the file.'}, status=500)

    if not extracted_text.strip():
        return JsonResponse({'error': 'Could not extract any text from the file.'}, status=400)

    # --- NLP Processing with Fine-Tuning for Name ---
    formatted_data = {}
    
    if nlp:
        # --- NAME HEURISTIC ---
        # 1. Look for a PERSON entity in the first ~300 characters of the resume
        first_part_of_resume = extracted_text[:300]
        doc_first_part = nlp(first_part_of_resume)
        names = [ent.text for ent in doc_first_part.ents if ent.label_ == 'PERSON']
        
        if names:
            # Often the first person found at the top is the correct one
            formatted_data['name'] = names[0]
        else:
            # Fallback to searching the whole document if no name is at the top
            doc_full = nlp(extracted_text)
            full_doc_names = [ent.text for ent in doc_full.ents if ent.label_ == 'PERSON']
            if full_doc_names:
                formatted_data['name'] = full_doc_names[0]

        # --- Other fields ---
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', extracted_text)
        if emails: formatted_data['email'] = emails[0]
        phones = re.findall(r'\(?\b\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', extracted_text)
        if phones: formatted_data['mobile_number'] = re.sub(r'\D', '', phones[0])

    # Use custom model for skills
    if custom_nlp:
        custom_doc = custom_nlp(extracted_text)
        found_skills = sorted(list({ent.text for ent in custom_doc.ents if ent.label_ == "SKILL"}))
        if found_skills:
            formatted_data['skills'] = ", ".join(found_skills)
    
    if not formatted_data:
        return JsonResponse({'error': 'Could not extract relevant information.'}, status=400)

    return JsonResponse({'success': True, 'data': formatted_data})


@login_required
@require_POST
def get_freelancer_ats_view(request, job_id):
    try:
        job = get_object_or_404(Job, pk=job_id)
        freelancer = get_object_or_404(FreelancerData, user=request.user)

        if not freelancer.resume or not freelancer.resume.path:
            return JsonResponse({'error': 'You must have a resume uploaded to check the score.'}, status=400)

        job_text = f"{job.title} {job.description} {' '.join([skill.name for skill in job.required_skills.all()])}"
        
        resume_path = freelancer.resume.path
        file_extension = os.path.splitext(resume_path)[1].lower()
        resume_text = ""

        if file_extension == '.pdf' and extract_pdf_text:
            resume_text = extract_pdf_text(resume_path)
        elif file_extension == '.docx' and docx:
            resume_text = '\n'.join([p.text for p in docx.Document(resume_path).paragraphs])
        elif file_extension in ['.jpg', '.jpeg', '.png'] and Image and pytesseract:
            resume_text = pytesseract.image_to_string(Image.open(resume_path))
        
        if not resume_text.strip():
            return JsonResponse({'error': 'Could not extract text from your resume.'}, status=400)

        score = get_resume_ats_score(job_text, resume_text)
        return JsonResponse({'success': True, 'score': score})

    except FileNotFoundError:
        return JsonResponse({'error': 'Your resume file could not be found. Please re-upload it.'}, status=404)
    except Exception as e:
        print(f"Error in ATS score calculation: {e}")
        return JsonResponse({'error': 'An unexpected server error occurred.'}, status=500)